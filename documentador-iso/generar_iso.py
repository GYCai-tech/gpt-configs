"""
generar_iso.py — Script autocontenido para generar fichas de procedimiento ISO
GÓMEZ Y CRESPO S.A. · ISO 9001/14001
Requiere: python-docx, PLANTILLA_PROCEDIMIENTO.docx en el mismo directorio

Uso desde Code Interpreter:
    import generar_iso
    ruta = generar_iso.generar(data)   # data es un dict con los campos del procedimiento
"""

import json
import os
import re
import sys
import tempfile

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Busca la plantilla en el mismo directorio que este script
HERE     = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(HERE, "PLANTILLA_PROCEDIMIENTO.docx")

AZUL  = "95B3D7"
VERDE = "E9EFB1"


# ── Utilidades XML ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_table_borders(table, sz=4):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)


def set_spacing(para, before=0, after=60):
    pPr     = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    old = pPr.find(qn("w:spacing"))
    if old is not None:
        pPr.remove(old)
    pPr.append(spacing)


def set_align(para, align):
    pPr = para._p.get_or_add_pPr()
    jc  = OxmlElement("w:jc")
    align_map = {
        WD_ALIGN_PARAGRAPH.LEFT:    "left",
        WD_ALIGN_PARAGRAPH.CENTER:  "center",
        WD_ALIGN_PARAGRAPH.RIGHT:   "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
    }
    jc.set(qn("w:val"), align_map.get(align, "left"))
    old = pPr.find(qn("w:jc"))
    if old is not None:
        pPr.remove(old)
    pPr.append(jc)


def add_run(para, text, size_pt=12, bold=False, italic=False,
            color_hex=None, font="Verdana"):
    run = para.add_run(text)
    run.font.name   = font
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return run


def add_field(para, field_type, size_pt=10):
    sz_val = str(int(size_pt * 2))

    def make_rPr():
        rPr = OxmlElement("w:rPr")
        sz  = OxmlElement("w:sz")
        sz.set(qn("w:val"), sz_val)
        rPr.append(sz)
        return rPr

    r_begin = OxmlElement("w:r")
    r_begin.append(make_rPr())
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc)
    para._p.append(r_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_type} "
    r_instr.append(instr)
    para._p.append(r_instr)

    r_end = OxmlElement("w:r")
    r_end.append(make_rPr())
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r_end.append(fc2)
    para._p.append(r_end)


def add_section_title(doc, text):
    p = doc.add_paragraph()
    add_run(p, text, size_pt=12, bold=True)
    set_spacing(p, before=120, after=60)
    return p


def blank(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=40)


# ── Header / Footer ────────────────────────────────────────────────────────────

def _set_wt(wt_node, text):
    wt_node.text = text
    if text != text.strip():
        wt_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def update_header(doc, data):
    hdr = doc.sections[0].header
    if not hdr.tables:
        return
    tbl = hdr.tables[0]

    cell_tit = tbl.cell(0, 1)
    p_tit = cell_tit.paragraphs[0]
    for r in list(p_tit._p.findall(qn("w:r"))):
        p_tit._p.remove(r)
    run = p_tit.add_run(f"{data['codigo']}: {data['nombre']}")
    run.font.name = "Verdana"
    run.font.bold = True
    run.font.size = Pt(14)

    cell_elab = tbl.cell(2, 0)
    wt_nodes = cell_elab._tc.findall(".//" + qn("w:t"))
    if wt_nodes:
        _set_wt(wt_nodes[0], f"Elaborado: {data['elaborado_por']}")

    seen, unique = set(), []
    for c in tbl.rows[2].cells:
        if id(c._tc) not in seen:
            seen.add(id(c._tc))
            unique.append(c)
    cell_apr = unique[-1]
    wt_nodes = cell_apr._tc.findall(".//" + qn("w:t"))
    if wt_nodes:
        _set_wt(wt_nodes[0], f"Revisado y Aprobado: {data['aprobado_por']}")


def update_footer(doc, data):
    ftr = doc.sections[0].footer
    if not ftr.tables:
        return
    tbl = ftr.tables[0]

    wt_nodes = tbl.cell(0, 0)._tc.findall(".//" + qn("w:t"))
    if wt_nodes:
        _set_wt(wt_nodes[0], f"{data['codigo']}: {data['nombre']}")

    wt_nodes = tbl.cell(0, 1)._tc.findall(".//" + qn("w:t"))
    if wt_nodes:
        _set_wt(wt_nodes[0], f"Fecha: {data['fecha']}")
        for wt in wt_nodes[1:]:
            r_elem = wt.getparent()
            if r_elem is not None and r_elem.getparent() is not None:
                r_elem.getparent().remove(r_elem)

    wt_nodes = tbl.cell(1, 0)._tc.findall(".//" + qn("w:t"))
    if wt_nodes:
        _set_wt(wt_nodes[0], f"Rev: {data['revision']}")
        for wt in wt_nodes[1:]:
            r_elem = wt.getparent()
            if r_elem is not None and r_elem.getparent() is not None:
                r_elem.getparent().remove(r_elem)

    cell_11 = tbl.cell(1, 1)
    p_elem  = cell_11.paragraphs[0]._p
    for r in list(p_elem.findall(qn("w:r"))):
        p_elem.remove(r)

    def _make_run(text):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Verdana")
        fonts.set(qn("w:hAnsi"), "Verdana")
        rPr.append(fonts)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t)
        return r

    p_elem.append(_make_run("Página "))
    add_field(cell_11.paragraphs[0], "PAGE",     size_pt=10)
    p_elem.append(_make_run(" de "))
    add_field(cell_11.paragraphs[0], "NUMPAGES", size_pt=10)


# ── Secciones del cuerpo ───────────────────────────────────────────────────────

def add_tabla_revisiones(doc, data):
    add_section_title(doc, "CONTROL DE REVISIONES")
    historial   = data["historial"]
    total_filas = len(historial) + 3 + 1
    tbl = doc.add_table(rows=total_filas, cols=5)
    tbl.style = "Table Grid"
    set_table_borders(tbl)
    for i, w in enumerate([Cm(1.332), Cm(1.669), Cm(7.752), Cm(2.588), Cm(2.912)]):
        for cell in tbl.columns[i].cells:
            cell.width = w
    for ri, entry in enumerate(historial):
        for ci, val in enumerate([entry["rev"], entry["fecha"], entry["descripcion"],
                                   entry.get("revisado", ""), entry.get("elaborado", "")]):
            c = tbl.cell(ri, ci)
            set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER if ci < 2 else WD_ALIGN_PARAGRAPH.LEFT)
            add_run(c.paragraphs[0], val, size_pt=11)
    for ci, h in enumerate(["REV", "FECHA", "DESCRIPCIÓN DE LOS CAMBIOS", "REVISADO Y APROBADO", "ELABORADO"]):
        c = tbl.cell(total_filas - 1, ci)
        set_cell_bg(c, AZUL)
        set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(c.paragraphs[0], h, size_pt=11)
    blank(doc)


def add_tabla_metadatos(doc, data):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    set_table_borders(tbl)
    for i, w in enumerate([Cm(5.419), Cm(5.417), Cm(5.417)]):
        tbl.columns[i].cells[0].width = w
    for ci, (label, val, color) in enumerate([
        ("FECHA:", data["fecha"], AZUL),
        ("REVISIÓN:", data["revision"], VERDE),
        ("PÁGINAS:", str(data["paginas"]), AZUL),
    ]):
        c = tbl.cell(0, ci)
        set_cell_bg(c, color)
        set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(c.paragraphs[0], label + " ", bold=True, size_pt=12)
        add_run(c.paragraphs[0], val, size_pt=12)
    blank(doc)


def add_indice(doc, data):
    # Secciones fijas con numeración ISO
    secciones = [
        (1, "Objeto.", False, []),
        (2, "Alcance.", False, []),
        (3, "Definiciones y Abreviaturas.", False, []),
        (4, "Responsabilidades.", False, []),
        (5, "Entradas y Salidas del Proceso.", False, []),
        (6, "Desarrollo.", True, data.get("desarrollo", [])),
        (7, "Archivo.", False, []),
        (8, "Diagrama de Flujo.", False, []),
        (9, "Referencias.", False, []),
        (10, "Anexos.", False, []),
    ]
    add_section_title(doc, "ÍNDICE")
    for num, titulo, tiene_sub, subitems in secciones:
        p = doc.add_paragraph()
        add_run(p, f"{num}. {titulo}", bold=True, size_pt=10, color_hex=AZUL)
        set_spacing(p, before=40, after=40)
        if tiene_sub:
            for item in subitems:
                p_sub = doc.add_paragraph()
                add_run(p_sub, f"    {item['num']} {item['titulo']}.", size_pt=10, color_hex=AZUL)
                set_spacing(p_sub, before=20, after=20)
    blank(doc)


def _add_multipar(doc, text):
    """Renderiza un campo de texto dividiendo por \\n\\n en párrafos separados."""
    parrafos = [pr.strip() for pr in text.replace("\r\n", "\n").split("\n\n") if pr.strip()]
    if not parrafos:
        parrafos = [text]
    for i, texto in enumerate(parrafos):
        p = doc.add_paragraph()
        add_run(p, texto, size_pt=12)
        set_align(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_spacing(p, before=0, after=60 if i < len(parrafos) - 1 else 80)


def add_objeto(doc, data):
    add_section_title(doc, "1. OBJETO")
    _add_multipar(doc, data["objeto"])
    blank(doc)


def add_alcance(doc, data):
    add_section_title(doc, "2. ALCANCE")
    _add_multipar(doc, data["alcance"])
    blank(doc)


def add_definiciones(doc, data):
    """Sección 3 — Definiciones y Abreviaturas (cláusula 3, ISO 9001:2015)"""
    add_section_title(doc, "3. DEFINICIONES Y ABREVIATURAS")
    definiciones = data.get("definiciones", [])
    if not definiciones:
        p = doc.add_paragraph()
        add_run(p, "No aplica.", italic=True, size_pt=12)
        blank(doc)
        return
    tbl = doc.add_table(rows=1 + len(definiciones), cols=2)
    tbl.style = "Table Grid"
    set_table_borders(tbl)
    for i, w in enumerate([Cm(5.0), Cm(11.253)]):
        for cell in tbl.columns[i].cells:
            cell.width = w
    for ci, h in enumerate(["TÉRMINO / ABREVIATURA", "DEFINICIÓN"]):
        c = tbl.cell(0, ci)
        set_cell_bg(c, AZUL)
        set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(c.paragraphs[0], h, bold=True, size_pt=11)
    for ri, item in enumerate(definiciones, 1):
        c0 = tbl.cell(ri, 0)
        add_run(c0.paragraphs[0], item.get("termino", ""), bold=True, size_pt=11)
        c1 = tbl.cell(ri, 1)
        add_run(c1.paragraphs[0], item.get("definicion", ""), size_pt=11)
    blank(doc)


def add_responsabilidades(doc, data):
    add_section_title(doc, "4. RESPONSABILIDADES")
    for rol in data.get("responsabilidades", []):
        p = doc.add_paragraph()
        add_run(p, rol["cargo"], bold=True, size_pt=12)
        set_align(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_spacing(p, before=80, after=20)
        for tarea in rol.get("tareas", []):
            p_t = doc.add_paragraph()
            add_run(p_t, f"• {tarea}", size_pt=12)
            set_spacing(p_t, before=0, after=40)
    blank(doc)


def add_entradas_salidas(doc, data):
    """Sección 5 — Entradas y Salidas del proceso (enfoque basado en procesos, ISO 9001:2015)"""
    add_section_title(doc, "5. ENTRADAS Y SALIDAS DEL PROCESO")
    entradas = data.get("entradas", [])
    salidas  = data.get("salidas", [])
    max_filas = max(len(entradas), len(salidas), 1)
    tbl = doc.add_table(rows=1 + max_filas, cols=2)
    tbl.style = "Table Grid"
    set_table_borders(tbl)
    for i, w in enumerate([Cm(8.126), Cm(8.127)]):
        for cell in tbl.columns[i].cells:
            cell.width = w
    for ci, h in enumerate(["ENTRADAS", "SALIDAS"]):
        c = tbl.cell(0, ci)
        set_cell_bg(c, AZUL)
        set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(c.paragraphs[0], h, bold=True, size_pt=11)
    for ri in range(max_filas):
        entrada = entradas[ri] if ri < len(entradas) else ""
        salida  = salidas[ri]  if ri < len(salidas)  else ""
        c0 = tbl.cell(ri + 1, 0)
        set_cell_bg(c0, VERDE)
        add_run(c0.paragraphs[0], f"• {entrada}" if entrada else "", size_pt=11)
        c1 = tbl.cell(ri + 1, 1)
        set_cell_bg(c1, VERDE)
        add_run(c1.paragraphs[0], f"• {salida}" if salida else "", size_pt=11)
    blank(doc)


def add_desarrollo(doc, data):
    add_section_title(doc, "6. DESARROLLO")
    for item in data.get("desarrollo", []):
        p = doc.add_paragraph()
        add_run(p, f"{item['num']}  {item['titulo']}", bold=True, size_pt=12)
        set_align(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
        set_spacing(p, before=120, after=60)
        # Soporta múltiples párrafos separados por \n\n o \n
        parrafos = [pr.strip() for pr in item["descripcion"].replace("\r\n", "\n").split("\n\n") if pr.strip()]
        if not parrafos:
            parrafos = [item["descripcion"]]
        for i, texto in enumerate(parrafos):
            p2 = doc.add_paragraph()
            add_run(p2, texto, size_pt=12)
            set_align(p2, WD_ALIGN_PARAGRAPH.JUSTIFY)
            set_spacing(p2, before=0, after=60 if i < len(parrafos) - 1 else 80)
    blank(doc)


def add_riesgos(doc, data):
    """Sección 7 — Riesgos y Oportunidades (cláusula 6.1, ISO 9001:2015)"""
    add_section_title(doc, "7. RIESGOS Y OPORTUNIDADES")
    riesgos = data.get("riesgos", [])
    if not riesgos:
        p = doc.add_paragraph()
        add_run(p, "No se han identificado riesgos específicos para este procedimiento.", italic=True, size_pt=12)
        blank(doc)
        return
    tbl = doc.add_table(rows=1 + len(riesgos), cols=4)
    tbl.style = "Table Grid"
    set_table_borders(tbl)
    for i, w in enumerate([Cm(5.5), Cm(1.5), Cm(5.5), Cm(3.752)]):
        for cell in tbl.columns[i].cells:
            cell.width = w
    for ci, h in enumerate(["RIESGO / OPORTUNIDAD", "TIPO", "ACCIÓN PREVISTA", "RESPONSABLE"]):
        c = tbl.cell(0, ci)
        set_cell_bg(c, AZUL)
        set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(c.paragraphs[0], h, bold=True, size_pt=11)
    for ri, item in enumerate(riesgos, 1):
        tbl.cell(ri, 0).paragraphs[0]
        add_run(tbl.cell(ri, 0).paragraphs[0], item.get("riesgo", ""), size_pt=11)
        tipo_cell = tbl.cell(ri, 1)
        set_align(tipo_cell.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(tipo_cell.paragraphs[0], item.get("tipo", "R"), size_pt=11)
        add_run(tbl.cell(ri, 2).paragraphs[0], item.get("accion", ""), size_pt=11)
        add_run(tbl.cell(ri, 3).paragraphs[0], item.get("responsable", ""), size_pt=11)
    blank(doc)


def add_indicadores(doc, data):
    """Sección 8 — Indicadores del Proceso (cláusula 9.1, ISO 9001:2015)"""
    add_section_title(doc, "8. INDICADORES DEL PROCESO")
    indicadores = data.get("indicadores", [])
    if not indicadores:
        p = doc.add_paragraph()
        add_run(p, "No se han definido indicadores específicos para este procedimiento.", italic=True, size_pt=12)
        blank(doc)
        return
    tbl = doc.add_table(rows=1 + len(indicadores), cols=5)
    tbl.style = "Table Grid"
    set_table_borders(tbl)
    for i, w in enumerate([Cm(3.5), Cm(4.0), Cm(2.0), Cm(2.5), Cm(4.253)]):
        for cell in tbl.columns[i].cells:
            cell.width = w
    for ci, h in enumerate(["INDICADOR", "FÓRMULA / MÉTODO", "META", "FRECUENCIA", "RESPONSABLE"]):
        c = tbl.cell(0, ci)
        set_cell_bg(c, AZUL)
        set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(c.paragraphs[0], h, bold=True, size_pt=11)
    for ri, item in enumerate(indicadores, 1):
        for ci, key in enumerate(["indicador", "formula", "meta", "frecuencia", "responsable"]):
            c = tbl.cell(ri, ci)
            set_cell_bg(c, VERDE)
            add_run(c.paragraphs[0], item.get(key, ""), size_pt=11)
    blank(doc)


def add_archivo(doc, data):
    """Sección 7 — Archivo (cláusula 7.5.3, ISO 9001:2015 — incluye plazo de conservación)"""
    add_section_title(doc, "7. ARCHIVO")
    filas = data.get("archivo", [])
    tbl   = doc.add_table(rows=1 + len(filas), cols=4)
    tbl.style = "Table Grid"
    set_table_borders(tbl)
    for i, w in enumerate([Cm(5.5), Cm(3.5), Cm(3.5), Cm(3.752)]):
        for cell in tbl.columns[i].cells:
            cell.width = w
    for ci, h in enumerate(["Documento / Registro", "Responsable", "Lugar", "Plazo de Conservación"]):
        c = tbl.cell(0, ci)
        set_cell_bg(c, AZUL)
        set_align(c.paragraphs[0], WD_ALIGN_PARAGRAPH.CENTER)
        add_run(c.paragraphs[0], h, bold=True, size_pt=11)
    for ri, fila in enumerate(filas, 1):
        for ci, key in enumerate(["documento", "responsable", "lugar", "plazo"]):
            c = tbl.cell(ri, ci)
            set_cell_bg(c, VERDE)
            add_run(c.paragraphs[0], fila.get(key, ""), size_pt=11)
    blank(doc)


def add_diagrama(doc):
    add_section_title(doc, "8. DIAGRAMA DE FLUJO")
    p = doc.add_paragraph()
    add_run(p, "[Insertar diagrama de flujo del procedimiento]",
            italic=True, color_hex="808080", size_pt=12)
    set_spacing(p, before=60, after=60)
    blank(doc)


def add_referencias(doc, data):
    """Sección 9 — Referencias, separadas en normativas externas e internas"""
    add_section_title(doc, "9. REFERENCIAS")
    refs = data.get("referencias", {})

    # Compatibilidad con formato antiguo (lista plana)
    if isinstance(refs, list):
        refs = {"normativas": [], "internas": refs}

    normativas = refs.get("normativas", [])
    internas   = refs.get("internas", [])

    if normativas:
        p = doc.add_paragraph()
        add_run(p, "Referencias normativas externas", bold=True, size_pt=12)
        set_spacing(p, before=60, after=20)
        for ref in normativas:
            p_r = doc.add_paragraph()
            add_run(p_r, f"• {ref}", size_pt=12)
            set_spacing(p_r, before=0, after=30)

    if interns := internas:
        p = doc.add_paragraph()
        add_run(p, "Documentos internos relacionados", bold=True, size_pt=12)
        set_spacing(p, before=60, after=20)
        for ref in interns:
            p_r = doc.add_paragraph()
            add_run(p_r, f"• {ref}", size_pt=12)
            set_spacing(p_r, before=0, after=30)

    if not normativas and not internas:
        p = doc.add_paragraph()
        add_run(p, "No aplica.", italic=True, size_pt=12)

    blank(doc)


def add_anexos(doc, data):
    add_section_title(doc, "10. ANEXOS")
    anexos = data.get("anexos", [])
    if anexos:
        for anexo in anexos:
            p = doc.add_paragraph()
            add_run(p, anexo, size_pt=12)
            set_align(p, WD_ALIGN_PARAGRAPH.JUSTIFY)
            set_spacing(p, before=0, after=40)
    else:
        p = doc.add_paragraph()
        add_run(p, "No aplica.", italic=True, size_pt=12)


# ── Validación ─────────────────────────────────────────────────────────────────

CAMPOS_OBLIGATORIOS = [
    "codigo", "nombre", "fecha", "revision", "paginas",
    "elaborado_por", "aprobado_por", "historial",
    "objeto", "alcance", "responsabilidades", "desarrollo",
]

def validar(data: dict):
    faltantes = [c for c in CAMPOS_OBLIGATORIOS if not data.get(c)]
    if faltantes:
        raise ValueError(f"Faltan campos obligatorios en el dict: {', '.join(faltantes)}")
    for entry in data.get("historial", []):
        if not entry.get("elaborado"):
            raise ValueError(
                f"El historial de revisión '{entry.get('rev')}' no tiene el campo 'elaborado' "
                "(quién elaboró el documento). ISO 9001 requiere trazabilidad de autoría."
            )


# ── Función principal ──────────────────────────────────────────────────────────

def generar(data: dict, output_dir: str = None) -> str:
    """
    Genera el DOCX del procedimiento ISO a partir de un dict.
    Devuelve la ruta del archivo generado.
    """
    validar(data)

    doc  = Document(TEMPLATE)
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        body.remove(child)
    if sectPr is not None:
        body.append(sectPr)

    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(1.75)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.43)

    doc.styles["Normal"].font.name = "Verdana"
    doc.styles["Normal"].font.size = Pt(12)

    update_header(doc, data)
    update_footer(doc, data)
    add_tabla_revisiones(doc, data)
    add_tabla_metadatos(doc, data)
    add_indice(doc, data)
    add_objeto(doc, data)
    add_alcance(doc, data)
    add_definiciones(doc, data)
    add_responsabilidades(doc, data)
    add_entradas_salidas(doc, data)
    add_desarrollo(doc, data)
    add_archivo(doc, data)
    add_diagrama(doc)
    add_referencias(doc, data)
    add_anexos(doc, data)

    nombre_safe = re.sub(r"[^\w\-]", "_", data.get("nombre", "procedimiento"))[:30]
    filename    = f"{data.get('codigo', 'DOC')}_{nombre_safe}.docx"
    out_dir     = output_dir or HERE
    out_path    = os.path.join(out_dir, filename)
    doc.save(out_path)
    print(f"Documento generado: {out_path}")
    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python generar_iso.py <archivo.json>")
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        data = json.load(f)
    generar(data)
